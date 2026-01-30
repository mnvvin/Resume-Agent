# backend/resume_utils.py
import os
import re
from pdfminer.high_level import extract_text
import docx

def extract_text_from_pdf(path_or_bytes):
    if isinstance(path_or_bytes, (bytes, bytearray)):
        temp_path = "temp_resume.pdf"
        with open(temp_path, "wb") as f:
            f.write(path_or_bytes)
        text = extract_text(temp_path)
        os.remove(temp_path)
    else:
        text = extract_text(path_or_bytes)
    return text or ""

def extract_text_from_docx(path_or_bytes):
    if isinstance(path_or_bytes, (bytes, bytearray)):
        temp_path = "temp_resume.docx"
        with open(temp_path, "wb") as f:
            f.write(path_or_bytes)
        doc = docx.Document(temp_path)
        os.remove(temp_path)
    else:
        doc = docx.Document(path_or_bytes)
    full_text = []
    for para in doc.paragraphs:
        if para.text:
            full_text.append(para.text)
    return "\n".join(full_text)

def extract_text_from_file(file_path: str) -> str:
    file_lower = file_path.lower()
    if file_lower.endswith(".pdf"):
        return extract_text_from_pdf(file_path)
    elif file_lower.endswith(".docx"):
        return extract_text_from_docx(file_path)
    else:
        raise ValueError("Unsupported file type")

def basic_resume_analysis(text: str):
    """
    Simple rule-based resume suggestions (no spaCy).
    """
    suggestions = []
    if not text:
        return {
            "word_count": 0,
            "suggestions": ["Could not read any text from the resume. Try another file or check the file format."],
        }

    words = re.findall(r"\w+", text)
    wcount = len(words)

    # length checks
    if wcount < 200:
        suggestions.append("Resume seems short. Add more detail about responsibilities, projects, and results.")
    elif wcount > 1500:
        suggestions.append("Resume seems long. Aim for 1–2 pages and keep content concise.")

    # contact checks
    if not re.search(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}", text):
        suggestions.append("No email address found. Add a professional email.")
    if not re.search(r"\b\d{7,}\b", text):
        suggestions.append("No phone number found. Add a contact phone number.")

    # section checks
    lowered = text.lower()
    if "experience" not in lowered and "work history" not in lowered:
        suggestions.append("Add a 'Work Experience' or 'Professional Experience' section.")
    if "education" not in lowered:
        suggestions.append("Add an 'Education' section.")
    if "skills" not in lowered and "technical skills" not in lowered:
        suggestions.append("Add a 'Skills' section to highlight your tools and technologies.")

    # achievement hints
    if "responsible for" in lowered and "increased" not in lowered and "%" not in lowered:
        suggestions.append("Use action verbs and quantify results (e.g., 'Increased sales by 20%').")

    if not suggestions:
        suggestions.append("Resume looks good on basic checks. Tailor it to the job description for better results.")

    return {
        "word_count": wcount,
        "suggestions": suggestions,
    }
